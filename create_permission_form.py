#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create a new Document
doc = Document()

# Add title
title = doc.add_paragraph()
title_run = title.add_run("EMPLOYEE INFORMATION SHARING AUTHORIZATION")
title_run.bold = True
title_run.font.size = Pt(14)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add spacing
doc.add_paragraph()

# Add company name
company = doc.add_paragraph()
company_run = company.add_run("Vanzyverden USA")
company_run.bold = True
company_run.font.size = Pt(12)
company.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add spacing
doc.add_paragraph()

# Add authorization text
doc.add_paragraph(
    "I hereby authorize Vanzyverden USA to share the following personal information "
    "with authorized parties as necessary for business purposes:"
)

# Add spacing
doc.add_paragraph()

# Add checkboxes for information types
doc.add_paragraph("☐  Home Address")
doc.add_paragraph("☐  Home Phone Number")
doc.add_paragraph("☐  Cell Phone Number")

# Add spacing
doc.add_paragraph()
doc.add_paragraph()

# Add acknowledgment text
doc.add_paragraph(
    "I understand that this authorization will remain in effect until I revoke it in writing. "
    "I acknowledge that I have read and understand this authorization."
)

# Add spacing
doc.add_paragraph()
doc.add_paragraph()

# Add employee information section
doc.add_paragraph("Employee Name (Print): _________________________________________________")
doc.add_paragraph()

# Add signature line
doc.add_paragraph("Employee Signature: _____________________________________________________")
doc.add_paragraph()

# Add date line
doc.add_paragraph("Date: ___________________________________________________________________")

# Add spacing
doc.add_paragraph()
doc.add_paragraph()

# Add footer note
footer = doc.add_paragraph(
    "Please return this completed form to Human Resources."
)
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer.runs[0]
footer_run.italic = True
footer_run.font.size = Pt(10)

# Save the document
doc.save('Employee_Information_Sharing_Authorization.docx')
print("Document created successfully: Employee_Information_Sharing_Authorization.docx")
